"""
Pypercraft Class
"""
import openai
import threading
from langchain import PromptTemplate
from docx import Document
from pypercraft.prompts import GENERATE_TITLE_PROMPT, GENERATE_INTRODUCTION_PROMPT, \
    GENERATE_BODY_PROMPT, GENERATE_CONCLUSION_PROMPT


class Pypercraft:
    """
    Pypercraft class that constructs a full paper
    """

    def __init__(self, query, topic, num_pages, tone, api_key):
        openai.api_key = api_key
        self.user_query = query
        self.topic = topic
        self.num_pages = num_pages
        self.tone = tone
        self.paper = {
            "title": "",
            "introduction": "",
            "body": "",
            "conclusion": ""
        }
        self.word_distribution = self.distribute_words(num_pages)
        self.system_role = """You are a helpful AI assistant that specializes in
        writing articles and papers."""

    def distribute_words(self, num_pages):
        """
        Determines the average number of words to assign each section
        """
        words_per_page = 500
        total_words = num_pages * words_per_page

        # Divide the total words among the sections
        intro_ratio = 0.1  # Introduction ratio, 10% of the total words
        conclusion_ratio = 0.1  # Conclusion ratio, 10% of the total words

        intro_words = int(total_words * intro_ratio)
        conclusion_words = int(total_words * conclusion_ratio)
        body_words = total_words - intro_words - conclusion_words

        return {
            "intro_words": intro_words,
            "body_words": body_words,
            "conclusion_words": conclusion_words
        }

    def generate_title(self):
        """
        Generates the title of the document
        :return: str
        """

        prompt = PromptTemplate.from_template(GENERATE_TITLE_PROMPT).format(
            tone=self.tone,
            idea=self.user_query,
            topic=self.topic)

        completion = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system",
                 "content": self.system_role},
                {"role": "user", "content": prompt}
            ]
        )

        self.paper['title'] = completion.choices[0].message['content']

        return completion.choices[0].message['content']

    def generate_introduction(self):
        """
        Generates the introduction of the document
        :return: str
        """

        prompt = PromptTemplate.from_template(GENERATE_INTRODUCTION_PROMPT).format(
            tone=self.tone,
            idea=self.user_query,
            topic=self.topic,
            num_words=self.word_distribution['intro_words'])

        completion = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system",
                 "content": self.system_role},
                {"role": "user", "content": prompt}
            ]
        )

        self.paper['introduction'] = completion.choices[0].message['content']

        return completion.choices[0].message['content']

    def generate_body(self):
        """
        Generates the body of the document
        :return: str
        """

        prompt = PromptTemplate.from_template(GENERATE_BODY_PROMPT).format(
            introduction=self.paper['introduction'],
            tone=self.tone,
            idea=self.user_query,
            topic=self.topic,
            num_words=self.word_distribution['body_words'])

        completion = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system",
                 "content": self.system_role},
                {"role": "user", "content": prompt}
            ]
        )

        self.paper['body'] = completion.choices[0].message['content']

        return completion.choices[0].message['content']

    def generate_conclusion(self):
        """
        Generates the conclusion of the document
        :return: str
        """

        prompt = PromptTemplate.from_template(GENERATE_CONCLUSION_PROMPT).format(
            tone=self.tone,
            idea=self.user_query,
            topic=self.topic,
            num_words=self.word_distribution['conclusion_words'])

        completion = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system",
                 "content": self.system_role},
                {"role": "user", "content": prompt}
            ]
        )

        self.paper['conclusion'] = completion.choices[0].message['content']

        return completion.choices[0].message['content']

    def construct(self, parallel=False):
        """
        Queries the API to generate a document
        :return:
        """
        if parallel:

            # Create thread objects
            thread1 = threading.Thread(target=self.generate_title)
            thread2 = threading.Thread(target=self.generate_introduction)
            thread3 = threading.Thread(target=self.generate_body)
            thread4 = threading.Thread(target=self.generate_conclusion)

            # Start the threads
            thread1.start()
            thread2.start()
            thread3.start()
            thread4.start()

            # Wait for threads to finish
            thread1.join()
            thread2.join()
            thread3.join()
            thread4.join()

        else:
            self.generate_title()
            self.generate_introduction()
            self.generate_body()
            self.generate_conclusion()

        return self.paper

    def export_docx(self, file_path):
        """
        Exports file as docx
        """
        doc = Document()

        # Adding title, introduction, body, and conclusion to the DOCX document
        doc.add_heading(self.paper["title"], level=1)
        doc.add_paragraph(self.paper["introduction"])
        doc.add_paragraph(self.paper["body"])
        doc.add_paragraph(self.paper["conclusion"])

        # Save as DOCX
        doc.save(file_path)
